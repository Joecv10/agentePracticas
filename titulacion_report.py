# titulacion_report.py
from pathlib import Path
from datetime import datetime
import os
import json
import openai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# la función que acabamos de crear
from titulacion_parser import parse_titulacion_excel

MODEL = "gpt-4o-mini"
TEMPERATURE = 0.4
openai.api_key = os.getenv("OPENAI_API_KEY")


def _add_heading(doc, text, level=1, center=False):
    p = doc.add_heading(text, level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_paragraphs(doc, items, bullet=False):
    for it in items:
        p = doc.add_paragraph(it)
        if bullet:
            p.style = 'List Bullet'


def _add_table_from_df(doc, df, caption=None):
    if caption:
        doc.add_paragraph(caption, style="Heading 2")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row):
            row_cells[j].text = str(val)


def generate_report_titulacion(excel_path: str | Path, save_dir: str | Path = "reports") -> Path:
    """
    Crea el informe de titulación (Indicador-95 formato nuevo) y devuelve la ruta del DOCX.
    """
    table1, table2, texts = parse_titulacion_excel(excel_path)

    # ---------- LLM para intro / conclusiones / recomendaciones ----------
    prompt_system = "Eres un redactor académico, tono formal."
    prompt_user = f"""Redacta para un informe ejecutivo:
- INTRODUCCIÓN: máximo 120 palabras sobre la importancia de planificar, ejecutar, hacer seguimiento y evaluar la titulación.
- CONCLUSIONES: 3–5 ítems sobre los hallazgos del análisis.
- RECOMENDACIONES: 3–5 ítems.

Contexto (JSON-resumen):
{json.dumps({
        "n_carreras": len(table1),
        "porcentaje_promedio_aprob": round(table1['Porcentaje de aprobación [%]'].mean(), 2)
    }, indent=2)}
"""

    resp = openai.chat.completions.create(
        model=MODEL,
        temperature=TEMPERATURE,
        messages=[
            {"role": "system", "content": prompt_system},
            {"role": "user",   "content": prompt_user},
        ],
    ).choices[0].message.content

    intro_txt, concl_txt, recom_txt = resp.split("CONCLUSIONES")[0], \
        "CONCLUSIONES" + resp.split("CONCLUSIONES", 1)[1].split("RECOMENDACIONES")[0], \
        "RECOMENDACIONES" + \
        resp.split("RECOMENDACIONES", 1)[1]

    # ------------- construir DOCX -----------------
    doc = Document()
    _add_heading(doc, "Indicador 95", level=0, center=True)

    _add_heading(doc, "INTRODUCCIÓN", level=1)
    _add_paragraphs(doc, [intro_txt])

    _add_heading(doc, "ANÁLISIS DE DATOS", level=1)

    # Acciones planificación / ejecución / seguimiento
    _add_heading(doc, "Acciones destacadas de planificación", level=2)
    _add_paragraphs(doc, texts['planificacion'], bullet=True)

    _add_heading(doc, "Acciones destacadas de ejecución", level=2)
    _add_paragraphs(doc, texts['ejecucion'], bullet=True)

    _add_heading(doc, "Acciones destacadas de seguimiento", level=2)
    _add_paragraphs(doc, texts['seguimiento'], bullet=True)

    _add_heading(doc, "Principales propuestas de mejora", level=2)
    _add_paragraphs(doc, texts['mejoras'], bullet=True)

    # Tablas
    _add_table_from_df(
        doc, table1, caption="Tabla 1. Aprobación de Trabajo de Titulación")
    _add_table_from_df(
        doc, table2, caption="Tabla 2. Distribución por modalidad de titulación")

    # Conclusiones y recomendaciones
    _add_heading(doc, "CONCLUSIONES", level=1)
    _add_paragraphs(doc, [concl_txt], bullet=False)

    _add_heading(doc, "RECOMENDACIONES", level=1)
    _add_paragraphs(doc, [recom_txt], bullet=False)

    # ---------------- guardar ---------------------
    Path(save_dir).mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    path = Path(save_dir) / f"Informe_Titulacion_{stamp}.docx"
    doc.save(path)
    return path
