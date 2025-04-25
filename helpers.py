import pandas as pd
import numpy as np, re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def show_headers(block):
    for i, c in enumerate(block.columns):
        print(f"{i:2} → {c[:80]}")   # first 80 chars so it fits

def get_block(df: pd.DataFrame, numero: int) -> pd.DataFrame:
    """
    Return the sub‑DataFrame that contains all columns belonging to Indicador_<numero>.
    Assumes each indicador block is 27 adjacent columns.
    """
    tag = f"Indicador_{numero}"
    cols = df.columns.tolist()

    # find the FIRST column whose header starts with that tag
    try:
        start = next(i for i, c in enumerate(cols) if c.startswith(tag))
    except StopIteration:
        raise ValueError(f"{tag} not found in DataFrame")

    end = start + 27          # 27 columns per block
    return df.iloc[:, start:end]


def parse_row(block_row: pd.Series) -> dict:
    """
    Convert a single row of a 27‑column indicador block into a structured dict.
    Returns keys:
       'respuesta_principal', 'actividades', 'objetivos',
       'n_estudiantes', 'n_docentes', 'n_administrativos',
       'otros_participantes', 'resultados'
    Splits multi‑value fields by '-' or '–'.
    Converts numeric strings to int when possible.
    """
    col_names = block_row.index.tolist()

    # helper to split by dash and strip whitespace
    def split_list(value):
        if pd.isna(value):
            return []
        return [x.strip() for x in str(value).split('-') if x.strip()]

    data = {
        "respuesta_principal": str(block_row[col_names[0]]).strip(),
        "actividades": split_list(block_row[col_names[3]]),
        "objetivos":   split_list(block_row[col_names[4]]),
        "n_estudiantes": split_list(block_row[col_names[5]]),
        "n_docentes":    split_list(block_row[col_names[6]]),
        "n_administrativos": split_list(block_row[col_names[7]]),
        "otros_participantes": split_list(block_row[col_names[8]]),
        "resultados": split_list(block_row[col_names[9]])
    }
    # cast counts to int where possible
    for key in ("n_estudiantes", "n_docentes", "n_administrativos"):
        data[key] = [int(x) for x in data[key] if re.fullmatch(r"\d+", x)]

    return data


def _find_col(block_row, contains: str):
    """Return the value from the first column whose header contains <contains>
       AND does *not* start with 'Puntos' or 'Comentarios'."""
    for col in block_row.index:
        hdr = col.lower()
        if contains in hdr and not hdr.startswith(("puntos", "comentarios")):
            return block_row[col]
    return None   # not found / blank

def _split_list(val):
    if pd.isna(val) or str(val).strip() == "":
        return []
    return [x.strip() for x in str(val).split('-') if x.strip()]

def parse_row(block_row: pd.Series) -> dict:
    d = {
        "respuesta_principal": str(block_row.iloc[0]).strip(),   # always column 0
        "actividades": _split_list(_find_col(block_row, "actividades realizadas")),
        "objetivos":   _split_list(_find_col(block_row, "objetivo de la actividad")),
        "n_estudiantes": _split_list(_find_col(block_row, "número de estudiantes")),
        "n_docentes":    _split_list(_find_col(block_row, "número de docentes")),
        "n_administrativos": _split_list(_find_col(block_row, "número de administrativos")),
        "otros_participantes": _split_list(_find_col(block_row, "otros participantes")),
        "resultados": _split_list(_find_col(block_row, "resultados de la actividad"))
    }
    # Cast counts to int where appropriate
    for key in ("n_estudiantes", "n_docentes", "n_administrativos"):
        d[key] = [int(x) for x in d[key] if re.fullmatch(r"\d+", x)]
    return d


def _first_val(row, keyword):
    """
    Return the cell value whose column header **contains** <keyword>
    but does NOT start with 'Puntos' or 'Comentarios'.
    """
    keyword = keyword.lower()
    for col in row.index:
        hdr = col.lower()
        if keyword in hdr and not hdr.startswith(("puntos", "comentarios")):
            return row[col]
    return None          # if nothing found

def _split_dash(val):
    if pd.isna(val) or str(val).strip() == "":
        return []
    return [x.strip() for x in str(val).split('-') if x.strip()]

def parse_row(row: pd.Series) -> dict:
    """Turn one survey response (row of 27 cols) into a tidy dict."""
    data = {
        "respuesta_principal": str(row.iloc[0]).strip(),

        "actividades":           _split_dash(_first_val(row, "actividades realizadas")),
        "objetivos":             _split_dash(_first_val(row, "objetivo de la actividad")),
        "n_estudiantes":         _split_dash(_first_val(row, "número de estudiantes")),
        "n_docentes":            _split_dash(_first_val(row, "número de docentes")),
        "n_administrativos":     _split_dash(_first_val(row, "número de administrativos")),
        "otros_participantes":   _split_dash(_first_val(row, "otros participantes")),
        "resultados":            _split_dash(_first_val(row, "resultados de la actividad")),
    }

    # cast numeric strings to ints where it makes sense
    for key in ("n_estudiantes", "n_docentes", "n_administrativos"):
        data[key] = [int(x) for x in data[key] if re.fullmatch(r"\d+", x)]

    return data


def aggregate_rows(rows: list[dict]) -> dict:
    """
    Combine many parsed rows (one per carrera) into a single dict ready for the LLM.
    — String lists are concatenated (duplicates removed, order preserved).
    — Numeric lists are summed.
    """
    agg = {
        "actividades": [],
        "objetivos": [],
        "otros_participantes": [],
        "resultados": [],
        "n_estudiantes_total": 0,
        "n_docentes_total": 0,
        "n_administrativos_total": 0,
        "n_carreras": len(rows),
    }

    seen = {k: set() for k in ("actividades", "objetivos", "otros_participantes", "resultados")}

    for r in rows:
        for key in ("actividades", "objetivos", "otros_participantes", "resultados"):
            for item in r[key]:
                if item not in seen[key]:
                    seen[key].add(item)
                    agg[key].append(item)

        agg["n_estudiantes_total"]     += sum(r["n_estudiantes"])
        agg["n_docentes_total"]        += sum(r["n_docentes"])
        agg["n_administrativos_total"] += sum(r["n_administrativos"])

    return agg


def save_report_to_docx(text: str, filename: str):
    """
    Crea un .docx con estilos sencillos.
    - La primera línea que empiece con '# ' se convierte en Título (Heading 0).
    - Líneas que empiecen con '## ' serán Heading 1.
    - El resto se añade como párrafos normales.
    """
    doc = Document()

    for line in text.splitlines():
        if line.startswith("# "):              # título principal
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):           # subtítulos
            doc.add_heading(line[3:].strip(), level=1)
        elif line.strip() != "":
            doc.add_paragraph(line)

    doc.save(filename)
    print(f"✅  Informe guardado en {filename}")